"""Word (DOCX) rapor üreteci — python-docx ile A4 formatlı rapor."""
import re
from datetime import datetime
from typing import Optional

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def generate_docx(
    title: str,
    content: str,
    company_name: str = "SustainHub",
    standard: str = "TSRS",
    version: int = 1,
    reporting_year: Optional[int] = None,
) -> bytes:
    if not DOCX_AVAILABLE:
        return content.encode("utf-8")

    doc = Document()

    # Sayfa boyutu A4
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    _set_header(doc, company_name, standard)
    _set_footer(doc, version)

    # Kapak sayfası
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover.add_run(f"\n\n\n{standard} Sürdürülebilirlik Raporu")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"{company_name}\n{reporting_year or datetime.now().year} Raporlama Yılı\nv{version}").font.size = Pt(14)

    doc.add_page_break()

    _parse_and_add_content(doc, content)

    from io import BytesIO
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _set_header(doc: "Document", company_name: str, standard: str) -> None:
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"{company_name} — {standard} Raporu  ")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
    _add_border_bottom(p)


def _set_footer(doc: "Document", version: int) -> None:
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ts = datetime.now().strftime("%d.%m.%Y")
    run = p.add_run(f"SustainHub Platform — sustainhub.ai — v{version} — {ts}  ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x75, 0x75, 0x75)
    # Sayfa numarası
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld)
    run._r.append(instrText)
    run._r.append(fld2)


def _add_border_bottom(paragraph) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1B5E20")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _parse_and_add_content(doc: "Document", content: str) -> None:
    lines = content.split("\n")
    for line in lines:
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue

        if stripped.startswith("### "):
            p = doc.add_heading(stripped[4:], level=3)
            p.runs[0].font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
        elif stripped.startswith("## "):
            p = doc.add_heading(stripped[3:], level=2)
            p.runs[0].font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
        elif stripped.startswith("# "):
            p = doc.add_heading(stripped[2:], level=1)
            p.runs[0].font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(stripped[2:], style="List Bullet")
            p.runs[0].font.size = Pt(11)
        elif re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s", "", stripped)
            p = doc.add_paragraph(text, style="List Number")
            p.runs[0].font.size = Pt(11)
        else:
            p = doc.add_paragraph()
            _add_inline_formatted(p, stripped)


def _add_inline_formatted(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(11)
        else:
            run = paragraph.add_run(part)
            run.font.size = Pt(11)
